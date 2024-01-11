import os
from docx import Document as DocxDocument
from database import db, Document
from transformers import pipeline
from PyPDF2 import PdfReader
import docx

# Initialize the QA pipeline from Hugging Face's Transformers
qa_pipeline = pipeline("question-answering", model="deepset/roberta-base-squad2")

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using the updated PdfReader class."""
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
    return text

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    doc = DocxDocument(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def extract_text(file_path):
    """Determine the file type and extract text accordingly."""
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

def process_uploaded_file(file_path, filename):
    text_content = extract_text(file_path)  # Implement this function based on file type (PDF, DOCX, etc.)
    new_document = Document(filename=filename, content=text_content)
    db.session.add(new_document)
    db.session.commit()


def search_documents_for_answer(question, documents):
    """
    Search through the provided documents for an answer to the question.
    """
    for doc in documents:
        result = qa_pipeline(question=question, context=doc.content)
        if result['score'] > 0.75:  # Adjust the threshold as needed
            return result['answer']
    return None

def read_pdf(file_path):
    with open(file_path, "rb") as file:
        pdf_reader = PdfReader(file)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
    return text

def read_word(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_txt(file_path):
    with open(file_path, "r") as file:
        text = file.read()
    return text



